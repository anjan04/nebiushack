#!/usr/bin/env python3
"""Generate journey.docx — the full story of building AutoRobot."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            t.rows[r_idx + 1].cells[c_idx].text = str(val)
    return t

# ============================================================
# TITLE PAGE
# ============================================================
title = doc.add_heading('AutoRobot: The Journey', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('An AI Scientist That Teaches Humanoid Robots to Walk\nNebius.Build SF Hackathon — March 15, 2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 120)

doc.add_paragraph()
para('This document captures every decision, pivot, failure, and lesson learned while building AutoRobot — a system that uses an LLM to autonomously discover reward configurations for humanoid locomotion, trained on NVIDIA H200 GPUs via Nebius Cloud.')

doc.add_page_break()

# ============================================================
# 1. THE IDEA
# ============================================================
heading('1. The Idea')

para('We started by exploring legged robot locomotion repos like HIMLoco (ICLR 2024), which uses reinforcement learning in Isaac Gym to teach robots to walk. The Nebius.Build SF hackathon focused on robotics + AI inference optimization — a perfect match.')

heading('Three Ideas We Considered', 2)
table(
    ['Idea', 'Approach', 'Why We Chose/Rejected'],
    [
        ['Voice-Commanded Robot', 'LLM as task planner + VLA for actions', 'Impressive demo but needs real robot hardware'],
        ['Inference-Optimized VLA', 'Quantize OpenVLA, benchmark speedup', 'Too infrastructure-focused, less creative'],
        ['LLM Reward Designer (chosen)', 'Autoresearch loop + Eureka-style reward gen', 'Novel combination, strong narrative, GPU-intensive'],
    ]
)

doc.add_paragraph()
para('We chose to combine two ideas:')
bullet('Karpathy\'s autoresearch', ' — an AI agent that makes one change at a time, evaluates, keeps/reverts, and accumulates improvements via git commits')
bullet('Eureka (NVIDIA, ICLR 2024)', ' — using LLMs to generate reward functions for robot RL training')
para('The pitch: "An AI scientist that runs 100 experiments overnight to teach a robot to walk, with zero human reward engineering."')

# ============================================================
# 2. INITIAL ARCHITECTURE
# ============================================================
heading('2. Initial Architecture')

para('We designed a simple 5-file system:')
table(
    ['File', 'Purpose', 'Agent Can Modify?'],
    [
        ['prepare.py', 'Isaac Gym env wrapper, observations, evaluation', 'No'],
        ['reward.py', 'Reward function — the LLM\'s canvas', 'Yes'],
        ['train.py', 'PPO training with fixed 3-min time budget', 'No'],
        ['agent.py', 'Orchestration loop: LLM → train → evaluate → git commit', 'No'],
        ['visualize.py', 'Score evolution charts and summary stats', 'No'],
    ]
)
doc.add_paragraph()
para('The original plan was: Isaac Gym Preview 4 + Ant quadruped + PPO via rl_games. Each experiment would train for 3 minutes, and the LLM would iterate on the reward function. Every experiment would be a git commit — the git log IS the demo.')

# ============================================================
# 3. PIVOT 1: ISAAC GYM → MUJOCO
# ============================================================
heading('3. First Pivot: Isaac Gym → MuJoCo + SB3')

heading('What Went Wrong', 2)
bullet('Isaac Gym Preview 4 requires a manual tarball download', ' from NVIDIA\'s developer portal — can\'t automate it')
bullet('Isaac Gym is deprecated', ' — NVIDIA recommends Isaac Lab now')
bullet('Our Nebius VM has an H200 GPU', ' which lacks RT cores required by Isaac Sim/Lab')

heading('What We Did', 2)
para('Switched to MuJoCo Humanoid-v5 + Stable-Baselines3 (SB3). Both are pip-installable.')

heading('Critical Bug Found', 2)
para('Our custom reward function was only running during evaluation, not during training. SB3 uses the environment\'s built-in reward by default. We had to write a CustomRewardWrapper that overrides step() to inject our reward into the training loop. Without this fix, the entire autoresearch loop was meaningless — the LLM\'s reward changes had zero effect on training.')

heading('The Speed Problem', 2)
table(
    ['Setup', 'Steps/sec', 'Steps in 3 min'],
    [
        ['Isaac Gym (planned)', '50,000+', '~9M'],
        ['MuJoCo CPU (actual)', '1,500', '~270K'],
        ['Gap', '33x slower', 'Barely learns anything'],
    ]
)
para('At 1,500 fps, the humanoid could barely learn to stand in 3 minutes, let alone walk. This made the approach unviable for a hackathon demo.')

# ============================================================
# 4. THE SONIC PAPER
# ============================================================
heading('4. The SONIC Paper')

para('A teammate committed sonic_paper.pdf to the repo. SONIC (NVIDIA) builds a humanoid foundation model using motion tracking as the core training signal — exactly the problem we\'re solving from the opposite angle.')

heading('SONIC\'s Approach vs Ours', 2)
table(
    ['', 'SONIC', 'AutoRobot'],
    [
        ['Mocap data', '700 hours, 100M+ frames', '1 walking clip'],
        ['GPUs', '128 GPUs, 32,000 GPU-hours', '1 GPU, ~hours'],
        ['Reward engineering', 'None (motion tracking)', 'LLM-automated'],
        ['Architecture', 'Custom encoder-decoder + FSQ', 'Standard MLP'],
    ]
)

heading('What We Incorporated', 2)
bullet('Exponential kernels: ', 'exp(-||error||² / temperature) for all tracking rewards')
bullet('SONIC weight hierarchy: ', 'orientation=0.5, position=1.0, velocity=1.0, penalties=-0.1')
bullet('Temperature tuning: ', 'start loose (forgiving), tighten over iterations')

# ============================================================
# 5. GPU REALITY CHECK
# ============================================================
heading('5. The GPU Reality Check')

para('This was the most surprising discovery of the hackathon. Our Nebius VM has an NVIDIA H200 — one of the most powerful AI GPUs in the world. But it can\'t run Isaac Lab.')

heading('Why?', 2)
para('The H200 is a compute-only GPU — it has no RT (Ray Tracing) cores. Isaac Lab depends on Isaac Sim, which initializes its RTX rendering engine even in headless mode. No RT cores = Isaac Sim won\'t start = Isaac Lab won\'t run. This is purely a software architecture limitation, not a hardware one.')

heading('Options We Evaluated', 2)
table(
    ['Option', 'GPU Accelerated?', 'H200 Compatible?', 'Verdict'],
    [
        ['Isaac Lab Stable', 'Yes (millions fps)', 'NO (needs RT cores)', 'Dead end'],
        ['Isaac Lab Newton (beta)', 'Yes', 'Yes (designed for it)', 'Too risky — beta bugs, humanoid flies away'],
        ['GEAR-SONIC', 'N/A', 'Yes', 'Training code NOT released'],
        ['MJX + Brax', 'Yes (100K-300K fps)', 'Yes', 'Good but needs JAX rewrite'],
        ['LocoMuJoCo', 'Yes (via JAX/MJX)', 'Yes', 'Winner — pip install, LAFAN1 built-in'],
        ['SB3 + MuJoCo', 'No (CPU only)', 'Yes', 'Already working but 200x too slow'],
    ]
)

# ============================================================
# 6. PIVOT 2: LOCOMUJOCO
# ============================================================
heading('6. Second Pivot: LocoMuJoCo + LAFAN1')

para('LocoMuJoCo checked every box:')
bullet('pip install loco-mujoco', ' — clean install, no tarballs')
bullet('LAFAN1 mocap data pre-retargeted', ' to Unitree H1 — auto-downloads from HuggingFace')
bullet('Built-in DeepMimic reward (MimicReward)', ' with 14 tunable weight parameters')
bullet('GPU-accelerated via JAX/MJX', ' — thousands of parallel environments')
bullet('Training a walking humanoid in ~36 minutes', ' on a single GPU')

heading('The Key Insight', 2)
para('Instead of having the LLM write arbitrary Python reward functions, we have it optimize the 14 weight parameters of LocoMuJoCo\'s built-in MimicReward. This is simpler, more robust, and leverages a proven reward architecture. The LLM outputs JSON instead of code.')

# ============================================================
# 7. PIVOT 3: CODE → JSON
# ============================================================
heading('7. Third Pivot: Python Code → JSON Weights')

table(
    ['', 'v1 (Python Code)', 'v3 (JSON Weights)'],
    [
        ['LLM outputs', 'Full reward function', '14 weight parameters'],
        ['Validation', 'AST parsing, import checking', 'JSON schema check'],
        ['Integration', 'Custom reward wrapper, monkey-patching', 'Pass weights to MimicReward constructor'],
        ['Robustness', 'LLM code can crash training', 'Weights always valid (just different quality)'],
        ['File modified', 'reward.py', 'reward_weights.json'],
    ]
)

# ============================================================
# 8. DEBUGGING WAR STORIES
# ============================================================
heading('8. Debugging War Stories')

heading('API Key Overwritten by Git Pull', 2)
para('We committed a placeholder .env to git. Every git pull overwrote the real API key with "your-nebius-api-key-here". The agent would run 5 experiments successfully, then silently fail on all subsequent calls. Lesson: never commit .env files, or use a different mechanism for secrets.')

heading('Thinking Models Return None', 2)
para('Qwen3-235B-Thinking returns content=None (the response is in a separate reasoning field). Our code assumed content was always a string. Every LLM call crashed with TypeError. Fix: content = response.choices[0].message.content or ""')

heading('Model Names Don\'t Match Documentation', 2)
para('Nebius documentation said "Qwen/Qwen3-235B-Thinking" but the actual API catalog had "Qwen/Qwen3-235B-A22B-Thinking-2507". We had to call models.list() to discover the correct names.')

heading('Python Output Buffering', 2)
para('We ran agent.py in the background with nohup. The output log was empty for hours. Python buffers stdout by default when not connected to a terminal. Fix: PYTHONUNBUFFERED=1 or python3 -u.')

heading('GPU Memory Leak (The Big One)', 2)
para('Every subprocess call to train_loco.py allocated ~108GB of H200 GPU memory via JAX. When the subprocess exited, JAX sometimes didn\'t release the memory. After 2-3 experiments, the GPU was full and new training processes couldn\'t start — causing "no supported devices found for platform CUDA" errors.')
para('Root cause: JAX pre-allocates the entire GPU memory by default.')
para('Fix: XLA_PYTHON_CLIENT_PREALLOCATE=false and XLA_PYTHON_CLIENT_MEM_FRACTION=0.7')

heading('Conda Terms of Service', 2)
para('Fresh Miniconda installations now require accepting TOS before creating environments. Our setup script failed silently. Had to add: conda tos accept --override-channels --channel https://repo.anaconda.com/pkgs/main')

heading('SSH Key Mismatch', 2)
para('The existing VM was created by a teammate with a different SSH key. We couldn\'t SSH in. Had to create a new VM through the Nebius console (using Playwright browser automation from Claude Code!) and add our SSH key during creation.')

heading('Subprocess Timeout vs JIT Compilation', 2)
para('Our training timeout was time_budget + 30 seconds. But JAX\'s JIT compilation takes 30-60 seconds on the first call, plus 2 seconds for LAFAN1 retargeting. Every experiment timed out. Fix: increased buffer to 120 seconds.')

# ============================================================
# 9. RESULTS
# ============================================================
heading('9. Results')

heading('Score Progression', 2)
table(
    ['Experiment', 'Score', 'Change', 'Key Insight'],
    [
        ['Baseline', '55.2', '—', 'Default MimicReward weights'],
        ['Exp 2 (Run 1)', '63.7', '+15.5%', 'Lower rpos_w_exp: 80→60'],
        ['Exp 3 (Run 3)', '82.6', '+33.5%', 'Further loosening: rpos_w_exp→50'],
        ['Exp 5 (Run 3)', '91.3', '+10.5%', 'Fine-tuning orientation weights'],
        ['Exp 8 (Run 3)', '97.9', '+7.3%', 'Added velocity tracking (rvel_w_sum)'],
        ['Exp 12 (Run 4)', '109.7', '+12%', 'Continued optimization'],
    ]
)

doc.add_paragraph()
para('Total improvement: 55.2 → 109.7 — nearly doubled the baseline, fully autonomously.')

heading('Training Stats', 2)
table(
    ['Metric', 'Value'],
    [
        ['GPU', 'NVIDIA H200 (143 GB VRAM)'],
        ['Parallel environments', '2,048 (GPU via JAX/MJX)'],
        ['Steps per experiment', '~37M'],
        ['Training throughput', '222,000 fps'],
        ['Time per experiment', '~3 minutes'],
        ['LLM', 'DeepSeek-V3 on Nebius Token Factory'],
        ['Robot', 'Unitree H1 (19 DOF)'],
        ['Reference motion', 'LAFAN1 walk1_subject1 (10,451 frames)'],
        ['Reward architecture', 'DeepMimic MimicReward (14 tunable weights)'],
    ]
)

heading('Key Discovery', 2)
para('The LLM autonomously discovered that loosening tracking temperatures produces dramatically better walking. This is counterintuitive — you\'d expect tighter tracking to produce more accurate motion imitation. But looser temperatures give the robot more freedom to explore during RL training, leading to more robust gaits that score higher on the overall tracking metric. This mirrors findings in the SONIC paper about the importance of temperature tuning.')

# ============================================================
# 10. TECHNICAL DECISIONS
# ============================================================
heading('10. Technical Decisions & Tradeoffs')

heading('Sequential vs Parallel Experiments', 2)
para('Eureka runs 16 reward candidates in parallel and picks the best. We chose sequential (one at a time) because:')
bullet('Simpler code — ', 'no GPU multiplexing or process management')
bullet('Compound improvements — ', 'each step builds on the last')
bullet('Better narrative — ', 'the git log tells a linear story')
bullet('Single GPU friendly — ', 'works on 1x H200')

heading('3-Minute Training Budget', 2)
para('Balances experiment throughput vs learning quality. With 2,048 parallel envs at 222K fps, 3 minutes gives ~37M steps — enough for meaningful learning but fast enough for ~20 experiments per hour.')

heading('DeepSeek-V3 over Qwen3-Thinking', 2)
para('Thinking models return content=None. Non-thinking models are simpler to integrate and reliably return the JSON weights we need.')

heading('JSON Weights over Python Code', 2)
para('More constrained = more robust. The LLM can\'t crash training with invalid code. Every set of weights produces a valid reward — just some are better than others.')

heading('Unitree H1 over Generic MuJoCo Humanoid', 2)
para('Real robot model with pre-retargeted mocap data. The generic Gymnasium Humanoid-v5 has no reference motions available — retargeting LAFAN1 to its unique skeleton would take weeks.')

# ============================================================
# 11. WHAT WE'D DO DIFFERENTLY
# ============================================================
heading('11. What We\'d Do Differently')

bullet('Start with LocoMuJoCo from day 1. ', 'We wasted hours on Isaac Gym, MuJoCo CPU, and Isaac Lab Newton before finding the right tool.')
bullet('Never commit API keys or .env files. ', 'Use environment variables or a secrets manager.')
bullet('Test the full pipeline end-to-end before building features. ', 'We built all 5 files before testing any of them together.')
bullet('Set XLA_PYTHON_CLIENT_PREALLOCATE=false immediately. ', 'The GPU memory leak cost us hours of debugging.')
bullet('Use PYTHONUNBUFFERED=1 for all background processes. ', 'Silent failures are the worst kind.')
bullet('Research GPU compatibility BEFORE provisioning VMs. ', 'H200 vs RT cores would have been caught in 5 minutes of research.')
bullet('Use a real database for experiment tracking, not git commits. ', 'Git is great for the narrative but terrible for querying results.')

# ============================================================
# 12. ARCHITECTURE DIAGRAM
# ============================================================
heading('12. Final Architecture')

para('The final system has 4 main components:')
para('1. LLM (DeepSeek-V3 on Nebius Token Factory) — proposes new reward weight configurations as JSON')
para('2. Training (LocoMuJoCo PPOJax on Nebius H200) — trains Unitree H1 with LAFAN1 mocap reference for 3 minutes')
para('3. Evaluation — compares tracking score against best, keeps or reverts')
para('4. Version Control — every experiment is a git commit with the weight configuration')
para('')
para('Loop: LLM proposes weights → PPO trains 3 min → evaluate → keep/revert → git commit → repeat')

# ============================================================
# 13. NEXT STEPS
# ============================================================
heading('13. Next Steps')

bullet('Switch to Unitree G1 — ', 'LocoMuJoCo has G1 with LAFAN1 data. One config change.')
bullet('Export policy to real robot — ', 'Extract numpy weights from JAX checkpoint, deploy via unitree_sdk2_python at 500 Hz')
bullet('Add domain randomization — ', 'friction, mass, motor noise for sim-to-real transfer')
bullet('Try more mocap clips — ', 'LAFAN1 has running, dancing, obstacle traversal')
bullet('Scale to more experiments — ', 'run overnight for 100+ experiments')
bullet('Compare LLM models — ', 'does Claude/GPT-4 find better weights than DeepSeek-V3?')

# ============================================================
# SAVE
# ============================================================
doc.save('/Users/anjan/nebiushack/docs/journey.docx')
print('Saved docs/journey.docx')
