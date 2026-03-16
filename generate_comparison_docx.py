#!/usr/bin/env python3
"""Generate autorobot_vs_sonic.docx from the markdown content."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

def heading(text, level=1):
    doc.add_heading(text, level=level)

def para(text):
    doc.add_paragraph(text)

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True if t.rows[0].cells[i].paragraphs[0].runs else False
    for r_idx, row in enumerate(rows):
        for c_idx in range(min(len(row), len(headers))):
            t.rows[r_idx + 1].cells[c_idx].text = str(row[c_idx])

# Parse the markdown file and convert to docx
with open('docs/autorobot_vs_sonic.md') as f:
    content = f.read()

# Split into sections
sections = re.split(r'^## ', content, flags=re.MULTILINE)

# Title
title = doc.add_heading('AutoRobot vs SONIC: A Technical Comparison', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('A detailed analysis comparing AutoRobot (LLM-automated reward optimization on 1 GPU)\nwith SONIC (NVIDIA\'s foundation model trained on 128 GPUs)')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(100, 100, 120)
doc.add_page_break()

# Process each section
for section in sections[1:]:  # skip the preamble
    lines = section.strip().split('\n')
    section_title = lines[0].strip()
    heading(section_title, level=1)

    i = 1
    current_table_headers = None
    current_table_rows = []

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines and horizontal rules
        if not line or line.startswith('---'):
            # Flush any pending table
            if current_table_headers and current_table_rows:
                table(current_table_headers, current_table_rows)
                current_table_headers = None
                current_table_rows = []
            i += 1
            continue

        # Sub-heading (### )
        if line.startswith('### '):
            if current_table_headers and current_table_rows:
                table(current_table_headers, current_table_rows)
                current_table_headers = None
                current_table_rows = []
            heading(line[4:], level=2)
            i += 1
            continue

        # Table row
        if line.startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if all(c.replace('-', '').replace(':', '') == '' for c in cells):
                # This is a separator row, skip
                i += 1
                continue
            if current_table_headers is None:
                current_table_headers = cells
            else:
                current_table_rows.append(cells)
            i += 1
            continue

        # Flush pending table before non-table content
        if current_table_headers and current_table_rows:
            table(current_table_headers, current_table_rows)
            current_table_headers = None
            current_table_rows = []

        # Bullet point
        if line.startswith('- **') or line.startswith('- '):
            text = line[2:].strip()
            # Clean markdown bold
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            bullet(text)
            i += 1
            continue

        # Code block - skip
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            i += 1
            continue

        # Regular paragraph
        text = line
        # Clean markdown formatting
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        if text:
            para(text)
        i += 1

    # Flush any remaining table
    if current_table_headers and current_table_rows:
        table(current_table_headers, current_table_rows)

doc.save('docs/autorobot_vs_sonic.docx')
print('Saved docs/autorobot_vs_sonic.docx')
